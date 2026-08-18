"""Career Coach - resume vs Naukri job description skill-gap tool.

Run:  C:\\Users\\tejas\\career-coach\\START.bat
Open: http://127.0.0.1:5055
"""
import io
import json
import os
import re
import shutil
import sqlite3
import datetime as dt

from flask import Flask, request, redirect, url_for, render_template, flash

from skills_db import SKILLS, ROLE_LABELS

# Optional AI analysis layer (Gemini). Falls back to the rule-based engine if no key.
import llm as _llm
llm_analyse = _llm.analyse
llm_clearance = _llm.clearance_plan_from_ai

# Durable store: mirrors reports + resume into the app's GitHub repo so they
# survive Render's free-plan disk wipes. Fail-safe: no-op if GH_TOKEN unset.
import reports_store

BASE = os.path.dirname(os.path.abspath(__file__))
# On Render the writable spot is a temp dir; locally it's the project folder.
DATA = os.environ.get("CAREER_DATA_DIR", BASE)
UPLOADS = os.path.join(DATA, "uploads")
DB = os.path.join(DATA, "career.db")
# Plain-text mirror of the saved resume. The SQLite DB (career.db) lives on a
# disk that Render's free plan wipes on every restart, which is exactly why the
# app kept demanding "upload your resume again". The mirror survives restarts,
# so on boot we re-seed the DB from it -> your resume is remembered for good.
RESUME_MIRROR = os.path.join(DATA, "resume_mirror.txt")
os.makedirs(UPLOADS, exist_ok=True)

# PIN lock: only needed when the app is public on the internet.
APP_PIN = os.environ.get("APP_PIN", "").strip()

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "career-coach-local")
app.config["MAX_CONTENT_LENGTH"] = 40 * 1024 * 1024  # 40 MB


@app.route("/health")
def health():
    """Lightweight status check - reports whether screenshot OCR is available
    on this deployment (Tesseract installed or not). No auth needed."""
    return {"ok": True, "ocr": ocr_available(),
            "role_count": len(SKILLS)}


@app.route("/ai-status")
def ai_status_route():
    """Reports whether the AI layer is active and, if not, exactly why.
    Useful for debugging 'rule-based mode' without guessing."""
    from llm import ai_status
    return ai_status()


@app.route("/ai-test")
def ai_test():
    """Temporary diagnostic: runs the real llm.analyse() with a cross-field
    SQL JD vs a non-technical resume and dumps the raw model response."""
    import llm, traceback
    jd = ("We are hiring a Database Developer. Required: strong SQL, PL/SQL, "
          "Oracle, ETL pipelines, data modelling, performance tuning. 5+ years.")
    res = ("I am a Training Manager with 8 years in L&D. Skills: ADDIE, "
           "Kirkpatrick, TNI, facilitation, coaching.")
    try:
        prov = llm.provider()
        st = llm.ai_status()
        # raw call to see exactly what the model returns
        raw = None
        raw_err = None
        try:
            raw = llm._call(llm._PROMPT.format(resume_text=res, jd_text=jd))
        except Exception as e:
            raw_err = f"{type(e).__name__}: {e}"
        parsed = llm._extract_json(raw) if raw else None
        rep = llm.analyse(jd, res)
        return {"provider": prov, "ai_status": st,
                "raw_err": raw_err,
                "raw_len": len(raw) if raw else 0,
                "raw_head": (raw[:500] if raw else None),
                "parsed_ok": bool(parsed),
                "parsed_keys": list(parsed.keys()) if isinstance(parsed, dict) else None,
                "returned": "AI dict" if rep else "None (fell back to rule-based)",
                "role": rep.get("role_label") if rep else None,
                "match": rep.get("match_pct") if rep else None}
    except Exception as e:
        return {"provider": llm.provider(), "error": f"{type(e).__name__}: {e}",
                "trace": traceback.format_exc()[-800:]}


@app.route("/env-debug")
def env_debug():
    """Temporary: list env var NAMES containing KEY/API/AI/GROQ (values hidden)
    so we can see exactly what Render passed without leaking secrets."""
    import os
    names = [k for k in os.environ
             if any(s in k.upper() for s in ("KEY", "API", "AI_", "GROQ", "GEMINI", "OPENAI"))]
    return {"env_var_names_seen": sorted(names)}


@app.before_request
def require_pin():
    """When APP_PIN is set (cloud), demand it once per browser session."""
    from flask import session
    if not APP_PIN:
        return None
    if request.endpoint in ("login", "static", "health") or request.path.startswith("/static"):
        return None
    if session.get("ok"):
        return None
    return render_template("login.html"), 401


@app.route("/login", methods=["GET", "POST"])
def login():
    from flask import session
    if not APP_PIN:
        return redirect(url_for("home"))
    if request.method == "POST":
        if (request.form.get("pin") or "").strip() == APP_PIN:
            session["ok"] = True
            session.permanent = True
            return redirect(url_for("home"))
        flash("Wrong PIN.", "bad")
    return render_template("login.html")


# --------------------------------------------------------------- database
def db():
    c = sqlite3.connect(DB)
    c.row_factory = sqlite3.Row
    return c


def init_db():
    with db() as c:
        c.execute("""CREATE TABLE IF NOT EXISTS resumes(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT, filename TEXT, text TEXT, uploaded TEXT)""")
        c.execute("""CREATE TABLE IF NOT EXISTS prompts(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            resume_id INTEGER, name TEXT, text TEXT, saved TEXT)""")
        c.execute("""CREATE TABLE IF NOT EXISTS jd(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT, role TEXT, source TEXT, jd_text TEXT,
            report TEXT, created TEXT)""")
    # Re-seed the most recent resume from the mirror so a Render restart never
    # asks the user to upload again.
    if not get_resume() and os.path.exists(RESUME_MIRROR):
        try:
            payload = json.loads(open(RESUME_MIRROR, encoding="utf-8").read())
            with db() as c:
                c.execute("INSERT INTO resumes(name,filename,text,uploaded) VALUES(?,?,?,?)",
                          ("My resume", payload.get("filename"), payload.get("text"),
                           payload.get("uploaded")))
        except Exception:
            pass

    # COLD-BOOT RECOVERY: Render's free disk is wiped on restart, so if the DB
    # has no reports after a restart, re-seed from GitHub (the durable copy).
    # This is what keeps your past reports alive after the app is closed.
    with db() as c:
        jd_count = c.execute("SELECT COUNT(*) FROM jd").fetchone()[0]
        res_count = c.execute("SELECT COUNT(*) FROM resumes").fetchone()[0]
    if reports_store.enabled():
        if jd_count == 0:
            for doc in reports_store.load_all_reports():
                try:
                    rep = doc.get("report")
                    if not rep:
                        continue
                    with db() as c:
                        c.execute(
                            "INSERT INTO jd(id,title,role,source,jd_text,report,created) "
                            "VALUES(?,?,?,?,?,?,?)",
                            (doc.get("id"),
                             doc.get("title", "Pasted ChatGPT analysis"),
                             doc.get("role", rep.get("role", "")),
                             "restored-from-github",
                             doc.get("jd_text", "") or "",
                             json.dumps(rep),
                             doc.get("created", "")))
                    print(f"[init] re-seeded report #{doc.get('id')} from GitHub")
                except Exception as e:
                    print(f"[init] re-seed report failed: {e}")
        if res_count == 0:
            r = reports_store.load_resume()
            if r:
                try:
                    with db() as c:
                        c.execute("INSERT INTO resumes(name,filename,text,uploaded) VALUES(?,?,?,?)",
                                  (r.get("name", "My resume"), r.get("filename"),
                                   r.get("text"), r.get("uploaded")))
                    print("[init] re-seeded resume from GitHub")
                except Exception as e:
                    print(f"[init] re-seed resume failed: {e}")


def save_prompt(resume_id, name, text):
    """Upsert the latest saved prompt for a given resume id."""
    with db() as c:
        c.execute("DELETE FROM prompts WHERE resume_id=?", (resume_id,))
        c.execute("INSERT INTO prompts(resume_id,name,text,saved) VALUES(?,?,?,?)",
                  (resume_id, name, text,
                   dt.datetime.now().strftime("%d %b %Y, %I:%M %p")))


def get_prompt(resume_id):
    with db() as c:
        return c.execute("SELECT * FROM prompts WHERE resume_id=? ORDER BY id DESC LIMIT 1",
                         (resume_id,)).fetchone()


def delete_prompt(pid):
    with db() as c:
        c.execute("DELETE FROM prompts WHERE id=?", (pid,))


def set_resume(name, filename, text, uploaded):
    """Persist a resume (named) in BOTH the DB and the mirror file."""
    with db() as c:
        c.execute("INSERT INTO resumes(name,filename,text,uploaded) VALUES(?,?,?,?)",
                  (name, filename, text, uploaded))
    try:
        with open(RESUME_MIRROR, "w", encoding="utf-8") as f:
            f.write(json.dumps({"name": name, "filename": filename,
                                "text": text, "uploaded": uploaded}))
    except Exception:
        pass
    # Durable copy in GitHub so it survives Render restarts.
    if reports_store.enabled():
        reports_store.save_resume(name, filename, text, uploaded)


def get_resumes():
    with db() as c:
        return c.execute("SELECT id,name,filename,uploaded FROM resumes ORDER BY id DESC").fetchall()


def get_resume(resume_id=None):
    with db() as c:
        if resume_id:
            row = c.execute("SELECT * FROM resumes WHERE id=?", (resume_id,)).fetchone()
            if row:
                return row
        return c.execute("SELECT * FROM resumes ORDER BY id DESC LIMIT 1").fetchone()


def delete_resume(resume_id):
    with db() as c:
        c.execute("DELETE FROM resumes WHERE id=?", (resume_id,))
        c.execute("DELETE FROM prompts WHERE resume_id=?", (resume_id,))


# Create tables at import time so the app works under gunicorn/Render too
# (gunicorn imports app:app and never runs the __main__ block below).
init_db()


# --------------------------------------------------------------- text extraction
def ocr_available():
    import pytesseract
    for p in (shutil.which("tesseract"),
              "/usr/bin/tesseract",
              r"C:\Program Files\Tesseract-OCR\tesseract.exe",
              r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
              os.path.expanduser(r"~\AppData\Local\Programs\Tesseract-OCR\tesseract.exe")):
        if p and os.path.exists(p):
            pytesseract.pytesseract.tesseract_cmd = p
            return True
    return False


def extract_text(storage):
    """Return (text, note) from an uploaded file: pdf / docx / txt / image."""
    name = (storage.filename or "file").lower()
    raw = storage.read()
    if not raw:
        return "", "empty file"

    if name.endswith(".pdf"):
        import pymupdf
        doc = pymupdf.open(stream=raw, filetype="pdf")
        txt = "\n".join(p.get_text() for p in doc)
        if len(txt.strip()) > 40:
            return txt, "PDF text layer"
        # scanned pdf -> OCR each page
        if ocr_available():
            import pytesseract
            from PIL import Image
            out = []
            for p in doc:
                pix = p.get_pixmap(dpi=200)
                out.append(pytesseract.image_to_string(
                    Image.open(io.BytesIO(pix.tobytes("png")))))
            return "\n".join(out), "scanned PDF via OCR"
        return txt, "PDF had no text layer and OCR is not installed"

    if name.endswith((".docx",)):
        import docx
        d = docx.Document(io.BytesIO(raw))
        parts = [p.text for p in d.paragraphs]
        for t in d.tables:
            for row in t.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts), "Word document"

    if name.endswith((".txt", ".md")):
        return raw.decode("utf-8", "ignore"), "text file"

    if name.endswith((".heic", ".heif")):
        # iPhone photos are often HEIC; Pillow can't read them without pyheif.
        try:
            import pyheif
            from PIL import Image
            img = Image.open(io.BytesIO(pyheif.read(raw).data))
            img.load()
        except Exception:
            return "", ("This screenshot is a HEIC file (iPhone format) I can't read. "
                        "On your phone, convert it to JPG/PNG (open it, tap Share > Save as image, "
                        "or screenshot it again) and upload that. Or just paste the job text.")
    else:
        # Anything Pillow can open -> OCR it (covers png/jpg/jpeg/webp/bmp/gif/tif/tiff).
        try:
            from PIL import Image
            img = Image.open(io.BytesIO(raw))
            img.load()
        except Exception:
            return "", f"unsupported or unreadable image: {name}"

    if not ocr_available():
        return "", ("Tesseract OCR is not installed yet, so I cannot read text out of "
                    "screenshots. Paste the job text into the box instead.")
    import pytesseract
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")
    if max(img.size) < 1400:                      # upscale small phone shots
        f = 1400 / max(img.size)
        img = img.resize((int(img.width * f), int(img.height * f)))
    try:
        txt = pytesseract.image_to_string(img)
    except Exception as e:
        return "", f"screenshot OCR failed ({type(e).__name__}). Paste the job text instead."
    return txt, "screenshot via OCR"

    return "", f"unsupported file type: {name}"


# --------------------------------------------------------------- matching engine
def norm(t):
    t = (t or "").lower()
    t = re.sub(r"[^a-z0-9&+/. ]+", " ", t)
    return re.sub(r"\s+", " ", t)


def hits(alias, text):
    """Word-boundary alias search; short aliases must stand alone."""
    a = norm(alias)
    if not a:
        return False
    return re.search(r"(?<![a-z0-9])" + re.escape(a) + r"(?![a-z0-9])", text) is not None


def detect_role(text):
    """Title line carries most weight, then the body."""
    t = norm(text)
    head = norm("\n".join(text.strip().splitlines()[:3])) if text.strip() else ""
    pats = {
        "rsm": ["regional sales", "sales manager", "area sales", "territory",
                "distributor", "revenue target", "channel sales"],
        "training manager": ["training manager", "manager training", "training head",
                             "capability manager", "manager learning and development",
                             "assistant manager training", "deputy manager training"],
        "l&d": ["l d manager", "learning and development", "learning development",
                "organisational development", "l d specialist", "l d executive"],
        "trainer": ["trainer", "facilitator", "process trainer", "soft skill",
                    "voice and accent", "instructional designer"],
    }
    score = {}
    for role, words in pats.items():
        score[role] = sum(5 * head.count(w) + t.count(w) for w in words)
    best = max(score, key=score.get)
    return best if score[best] else "training manager"


def build_verdict(rep, resume_text):
    """Plain-English synthesis: what this JD means for THIS candidate specifically."""
    gaps = rep["gaps"]
    near = [g for g in gaps if g.get("near")]
    true_gaps = [g for g in gaps if not g.get("near")]
    have = rep["have"]
    role = rep["role_label"]

    lines = []
    if have:
        top = ", ".join(h["key"].split(" (")[0] for h in have[:3])
        lines.append(f"Your strongest cards for this {role} role are already on your resume: {top}.")
    if near:
        n = ", ".join(g["key"].split(" (")[0] for g in near[:4])
        lines.append(f"You're closer than the JD makes it look on: {n}. "
                     f"You already do the work - you just lack the framework name or the tool. "
                     f"That's a vocabulary-and-proof gap, not a from-zero gap, and it's the fastest to close.")
    if true_gaps:
        t = ", ".join(g["key"].split(" (")[0] for g in true_gaps[:4])
        lines.append(f"Genuine new skills to learn: {t}. "
                     f"These need real practice, not just words - use your own NHT programme as the lab.")
    if not gaps and have:
        lines.append("Nothing in this JD is missing from your background. Lead with proof and numbers in the interview.")
    if not have and gaps:
        lines.append("This role is a stretch from your current resume - build the adjacent proof first.")
    bottom = (f"Bottom line: {len(near)} skill(s) you can reframe from what you already do, "
              f"{len(true_gaps)} you need to actually learn. "
              f"Start with the 'near' ones - they turn into interview wins in days, not months.")
    lines.append(bottom)
    return " ".join(lines)


# --------------------------------------------------------------- routes
def analyse(jd_text, resume_text):
    jd, res = norm(jd_text), norm(resume_text)
    role = detect_role(jd_text)

    asked, gaps, have = [], [], []
    for s in SKILLS:
        in_jd = any(hits(a, jd) for a in s["aliases"]) or hits(s["key"], jd)
        if not in_jd:
            continue
        in_res = any(hits(a, res) for a in s["aliases"]) or hits(s["key"], res)
        # ADJACENT: resume shows related work even if it doesn't name the framework
        adj = s.get("adjacent", [])
        adj_hit = [a for a in adj if hits(a, res)]
        rec = dict(key=s["key"], why=s["why"], learn=s["learn"], proof=s["proof"],
                   link=s.get("link", ""), tools=s.get("tools", []),
                   matched=[a for a in s["aliases"] if hits(a, jd)][:6],
                   bridge=adj_hit[:3])
        asked.append(rec)
        if in_res:
            have.append(rec)
        else:
            # mark as 'near' if resume shows adjacent work -> reframe as vocabulary gap
            rec["near"] = bool(adj_hit)
            gaps.append(rec)

    # role-critical skills the JD implies even if not literally worded
    implied = []
    for s in SKILLS:
        if role in s["roles"] and s["key"] not in [a["key"] for a in asked]:
            if not (any(hits(a, res) for a in s["aliases"]) or hits(s["key"], res)):
                adj = s.get("adjacent", [])
                adj_hit = [a for a in adj if hits(a, res)]
                implied.append(dict(key=s["key"], why=s["why"], learn=s["learn"],
                                    proof=s["proof"], link=s.get("link", ""), tools=s.get("tools", []),
                                    matched=[], near=bool(adj_hit), bridge=adj_hit[:3]))

    asked_count = len(asked)
    # Honest match %: only meaningful when the JD actually shares our skill
    # vocabulary. If it matched NONE of the 21 known skills, the JD is from a
    # different field entirely -> do NOT show a flattering number.
    if asked_count == 0:
        match_pct = 0
        out_of_domain = True
    else:
        match_pct = round(100 * len(have) / asked_count)
        # never claim a confident match on a single weak signal
        if asked_count < 3 and match_pct >= 100:
            match_pct = 0
        out_of_domain = False
    total = asked_count or 1
    rep = dict(role=role, role_label=ROLE_LABELS[role],
               match_pct=match_pct,
               have=have, gaps=gaps, implied=implied[:6],
               asked_count=asked_count, out_of_domain=out_of_domain,
               generated=dt.datetime.now().strftime("%d %b %Y, %I:%M %p"))
    rep["verdict"] = build_verdict(rep, res)
    return rep


def clearance_plan(rep, jd_text, resume_text):
    """Build the point-by-point 'Interview Clearance Plan' the user asked for:
    Job / Experience required / Skills lacking in resume / Skills required /
    How to learn / Chances if learned / Resources (links, books, YouTube, free tools).
    This is more structured than the gap cards and reads like a checklist."""
    res = norm(resume_text)
    jd = norm(jd_text)
    role_label = rep["role_label"]

    # Roll up every skill the JD asks for but the resume is weak on (gaps + implied)
    items = []
    for g in (rep["gaps"] + rep["implied"]):
        # "how to learn" = the learn steps (already free-first)
        learn = g.get("learn", [])
        # "resources" = the link (official) + tools block (books/yt/tools/ai)
        tools = g.get("tools", [])
        # split tools into labelled buckets for clean pointers
        books, yt, free_tools, ai = [], [], [], []
        for t in tools:
            tl = t.lower()
            if "youtube" in tl:
                yt.append(t)
            elif "book" in tl or "read" in tl or "article" in tl or "atd" in tl \
                 or "kirkpatrick" in tl or "shrm" in tl or "iso" in tl or "nngroup" in tl:
                books.append(t)
            elif "ai:" in tl or "chatgpt" in tl or "gemini" in tl or "copilot" in tl:
                ai.append(t)
            else:
                free_tools.append(t)
        link = (g.get("link") or "").strip()
        items.append(dict(
            key=g["key"],
            why=g.get("why", ""),
            proof=g.get("proof", ""),
            near=bool(g.get("near")),
            bridge=g.get("bridge", []),
            learn=learn,
            link=link,
            books=books, yt=yt, free_tools=free_tools, ai=ai))
    return dict(role_label=role_label, role_read=ROLE_LABELS.get(rep["role"], role_label),
                asked=len(rep["have"]) + len(rep["gaps"]),
                have=len(rep["have"]), gaps=len(rep["gaps"]),
                items=items)


ROLE_EXPERIENCE = {
    "training manager": "Typically 4-7 yrs in training/L&D with team coordination; shows you can own a curriculum end-to-end and manage people, not just deliver sessions.",
    "l&d": "Typically 3-6 yrs in L&D / OD; shows you design learning strategy and run programmes, not only facilitate.",
    "trainer": "Typically 1-4 yrs facilitating; shows you can deliver and design sessions and handle a demo round cold.",
    "rsm": "Typically 4-8 yrs with a carried revenue/quota number and team/territory handling; L&D-to-RSM is a stretch without quota history.",
}
def interview_questions(rep, jd_text, resume_text):
    """Generate the questions THIS employer will most likely ask this candidate,
    drawn from the gaps the engine found plus the candidate's own resume.

    This is the part a plain ChatGPT paste can't do well: it knows the gaps and
    your real background, so every question comes with a tailored 'your answer'
    built from your resume and the bridge skill you already have."""
    res = norm(resume_text)

    def answer_for(rec):
        # Prefer the candidate's own resume wording that shows adjacent work.
        if rec.get("bridge"):
            return (f"Use the {rec['key']} angle: on your resume you already do "
                    f"'{rec['bridge'][0]}'. Reframe it as {rec['key']} — same work, "
                    f"framework name + one proof point. {rec.get('proof', '')}")
        if rec.get("near"):
            return (f"You're close on {rec['key']} — you do the work but may not name "
                    f"the framework. Say: '{rec.get('proof', '')}'")
        return (f"Be honest you're building it, and pivot to the closest thing you do: "
                f"{rec.get('proof', '')}")

    items = []
    for g in (rep["gaps"] + rep["implied"])[:6]:
        items.append(dict(
            q=f"The JD asks for {g['key']}. How do you handle that today, and where's your proof?",
            tip=g.get("why", ""),
            a=answer_for(g)))
    # A couple of role-level behavioural questions.
    items.append(dict(
        q=f"Why are you moving from training operations to a {rep['role_label']} role?",
        tip="They'll probe motivation. Tie it to ownership you already have.",
        a="Point to owning the 12-trainer team, the NHT/Technical curriculum and MIS "
          "governance — you already operate at manager scope, you want the title and the "
          "strategic remit that comes with it."))
    items.append(dict(
        q=f"What's your biggest gap for a {rep['role_label']} role, and how are you closing it?",
        tip="They reward self-awareness here, not 'I have no weaknesses'.",
        a="Name the top true gap from above, show the concrete learning step you've started "
          "(one of the 'how to learn it' links), and a date you'll have proof."))
    return items


# --------------------------------------------------------------- cumulative
# Maps each skill in the KB to one of the 12 "Core Career Skill Stack"
# categories from the user's cumulative-recruiter-demand spec.
CATEGORY_MAP = {
    "BPO / Operations": [],
    "Customer Service": [],
    "Sales": ["Sales target ownership & pipeline management",
              "Channel / distributor management",
              "CRM tools (Salesforce / Zoho / LeadSquared)"],
    "Training & Facilitation": ["Train-the-Trainer / facilitation certification",
                                "Communication & business English / presentation",
                                "Onboarding / new hire training programme design"],
    "Learning & Development": ["ADDIE instructional design",
                               "Kirkpatrick training evaluation (L1-L4)",
                               "TNI / TNA (Training Needs Identification)",
                               "LMS administration",
                               "e-learning authoring (Articulate / Rise / Captivate)",
                               "Coaching & feedback models (GROW, SBI)",
                               "Content / SOP documentation"],
    "Sales Training": [],
    "Leadership & People Management": ["Team leadership & performance management",
                                        "Stakeholder / business partnering"],
    "Performance Management": ["Team leadership & performance management"],
    "Data / Reporting / Analytics": ["Excel advanced (pivots, lookups, dashboards)",
                                     "Power BI / data visualisation",
                                     "Learning analytics / MIS reporting"],
    "Digital / AI / Technology": ["AI tools for L&D (ChatGPT/Gemini for content)"],
    "Communication & Stakeholder Management": ["Communication & business English / presentation",
                                               "Stakeholder / business partnering"],
    "Learning Technology / LMS": ["LMS administration",
                                  "e-learning authoring (Articulate / Rise / Captivate)"],
}

TIER_LABEL = {1: "TIER 1 — ESSENTIAL / MUST-HAVE",
              2: "TIER 2 — HIGH-VALUE SKILLS",
              3: "TIER 3 — DIFFERENTIATING SKILLS",
              4: "TIER 4 — LOW PRIORITY / ROLE-SPECIFIC"}

LEVEL_LABEL = {
    "demonstrated": "✅ DEMONSTRATED — clear evidence in my resume",
    "transferable": "🔶 TRANSFERABLE / ADJACENT — experience overlaps, keyword may differ",
    "partial": "🟡 PARTIAL GAP — related experience, missing part of the capability",
    "gap": "🔴 GENUINE GAP — no meaningful resume or transferable evidence",
    "na": "⚪ Not relevant to my target career",
    "unknown": "❔ Resume not saved — level unknown",
}

ACTION_LABEL = {
    "MARKET MORE STRONGLY": "YOU ALREADY HAVE THIS — MARKET IT MORE STRONGLY",
    "STRENGTHEN / FORMALIZE": "STRENGTHEN / FORMALIZE (you have adjacent experience)",
    "STRENGTHEN / LEARN": "STRENGTHEN / LEARN (partial — close the missing part)",
    "LEARN": "LEARN (genuine gap — recruiters want it, you lack it)",
    "NOT PRIORITY": "DO NOT PRIORITIZE YET (low demand / niche)",
}

# Canonical skill -> one of the 7 target career categories (spec #12). Maps the
# 22 KB skills; learned skills fall back to "Other".
CANON_CATEGORY = {
    "ADDIE instructional design": "Learning & Development",
    "Kirkpatrick training evaluation (L1-L4)": "Learning & Development",
    "TNI / TNA (Training Needs Identification)": "Learning & Development",
    "LMS administration": "Learning Technology / LMS",
    "e-learning authoring (Articulate / Rise / Captivate)": "Learning Technology / LMS",
    "Train-the-Trainer / facilitation certification": "Training & Facilitation",
    "Coaching & feedback models (GROW, SBI)": "Training & Facilitation",
    "Content / SOP documentation": "Training & Facilitation",
    "Excel advanced (pivots, lookups, dashboards)": "Data / Reporting / Analytics",
    "Power BI / data visualisation": "Data / Reporting / Analytics",
    "Learning analytics / MIS reporting": "Data / Reporting / Analytics",
    "Sales target ownership & pipeline management": "Sales",
    "Channel / distributor management": "Sales",
    "CRM tools (Salesforce / Zoho / LeadSquared)": "Sales",
    "Team leadership & performance management": "Leadership & People Management",
    "Stakeholder / business partnering": "Communication & Stakeholder Management",
    "Onboarding / new hire training programme design": "Training & Facilitation",
    "Compliance / audit readiness of training records": "Learning & Development",
    "Communication & business English / presentation": "Communication & Stakeholder Management",
    "AI tools for L&D (ChatGPT/Gemini for content)": "Digital / AI / Technology",
    "Vendor & training budget management": "Leadership & People Management",
    "BPO / Operations": "BPO / Operations",
    "Customer Service": "Customer Service",
    "Sales Training": "Sales Training",
}

# Inverted view: category -> list of canonical skill names (used by the Core
# Career Skill Stack). Built once from CANON_CATEGORY.
CATEGORY_MEMBERS = {}
for _sk, _cat in CANON_CATEGORY.items():
    CATEGORY_MEMBERS.setdefault(_cat, []).append(_sk)


# ============================================================== CUMULATIVE
# TRUE cumulative recruiter-demand engine.
#
# Design rules (spec 2026-08 upgrade):
#  * ONE underlying JD dataset feeds Top15, tiers, roadmap, career direction,
#    and the final block (RULE 6/7). No separate "learned skills" silo.
#  * A canonical skill's demand = (# of unique JDs that demand it / total
#    unique JDs analyzed) * 100, with the numerator and denominator ALWAYS
#    shown (RULE 5).
#  * A JD "demands" a skill if EITHER the pasted ChatGPT analysis flagged it
#    (have/gaps/implied — provenance from the JD) OR the JD text literally
#    names a canonical alias. Counted ONCE per JD (no double credit).
#  * Synonym phrases map to one canonical skill; original JD wording is kept
#    as audit evidence (RULE 4, spec #19).
#  * Resume evidence never assumed from a JD alone (RULE 10); four levels:
#    demonstrated / transferable / partial / genuine gap (spec #8/10).
def cumulative_analysis(exclude_id=None, resume_override=None):
    """Cumulative Recruiter-Demand Analysis over ALL saved JDs.

    resume_override lets a caller (validation harness) pass a resume text
    without touching the stored one. Returns one dict consumed by both the
    /cumulative page and the per-report block.
    """
    r = get_resume()
    resume_text = (resume_override if resume_override is not None
                   else (r["text"] if r else "")) or ""
    res = norm(resume_text)
    has_resume = bool(res.strip())

    with db() as c:
        rows = c.execute(
            "SELECT id,title,role,created,jd_text,report FROM jd ORDER BY id").fetchall()
    if exclude_id is not None:
        rows = [row for row in rows if row["id"] != exclude_id]

    jd_count = len(rows)
    small = jd_count < 5

    # ---- 1. Build the canonical universe: KB skills + "learned" new skills -
    # Each canonical entry: name, aliases, category, kind.
    canon = []  # list of dict(name, aliases, category, kind, why, key)
    kb_by_norm = {}
    for s in SKILLS:
        name = s["key"]
        aliases = [name] + list(s.get("aliases", []))
        canon.append(dict(name=name, aliases=aliases,
                          category=CANON_CATEGORY.get(name, "Other"),
                          kind="kb", why=s.get("why", ""),
                          adjacent=list(s.get("adjacent", []))))
        for a in aliases:
            kb_by_norm.setdefault(norm(a), name)

    # Discover "learned" canonical skills: requirement names ChatGPT pulled
    # from JDs that are NOT one of the 22 KB skills/aliases. Exact match only
    # (no substring) so "Salesforce CPQ" is never swallowed by "salesforce".
    _GENERIC = {"degree","bachelors","bachelor","master","mba","phd","experience",
                "years","education","certification","certifications","qualification",
                "qualifications","skills","ability","knowledge","responsibilities",
                "requirements","communication","presentation","training","learning"}
    learned_names = {}
    for row in rows:
        try:
            rep = (json.loads(row["report"]) if isinstance(row["report"], str)
                   else row["report"]) or {}
        except Exception:
            rep = {}
        for nm in (rep.get("required_skills") or []):
            nm = (nm or "").strip().rstrip(".").strip()
            if len(nm) < 3 or len(nm) > 60 or norm(nm) in _GENERIC:
                continue
            if norm(nm) in kb_by_norm:
                continue  # already a KB skill
            learned_names.setdefault(nm, 0)
            learned_names[nm] += 1
    for nm, n in sorted(learned_names.items(), key=lambda x: -x[1]):
        name = nm
        canon.append(dict(name=name, aliases=[name],
                          category=CANON_CATEGORY.get(name, "Other"),
                          kind="learned", why="", adjacent=[]))

    # ---- 2. Per-JD demand scan (provenance + required/preferred + phrases) -
    # demand[name] = dict(jids=set, required=set, preferred=set,
    #                      phrases=set, titles=list, jd_ids=list)
    demand = {c["name"]: dict(jids=set(), required=set(), preferred=set(),
                              phrases=set(), titles=[], jd_ids=[])
              for c in canon}

    def _classify_req(cat):
        if cat == "preferred":
            return "preferred"
        return "required"  # explicitly_required or implied -> counted as required demand

    for row in rows:
        jid = row["id"]
        jd_norm = norm(row["jd_text"] or "")
        try:
            rep = (json.loads(row["report"]) if isinstance(row["report"], str)
                   else row["report"]) or {}
        except Exception:
            rep = {}
        title = (row["title"] or "JD %d" % jid)
        # (a) provenance from the pasted ChatGPT analysis
        for g in (rep.get("gaps") or []):
            k = (g.get("key") or "").strip()
            if k in demand:
                d = demand[k]
                if jid not in d["jids"]:
                    d["jids"].add(jid); d["titles"].append(title); d["jd_ids"].append(jid)
                rc = _classify_req(g.get("category"))
                d[rc].add(jid)
                q = (g.get("jd_quote") or "").strip()
                if q:
                    d["phrases"].add(q)
        for h in (rep.get("have") or []):
            k = (h.get("key") or "").strip()
            if k in demand:
                d = demand[k]
                if jid not in d["jids"]:
                    d["jids"].add(jid); d["titles"].append(title); d["jd_ids"].append(jid)
        for im in (rep.get("implied") or []):
            k = (im.get("key") or "").strip()
            if k in demand:
                d = demand[k]
                if jid not in d["jids"]:
                    d["jids"].add(jid); d["titles"].append(title); d["jd_ids"].append(jid)
                d["required"].add(jid)
        # (b) literal alias scan — catches skills the ChatGPT analysis may have
        #     missed; credited once per JD.
        for c in canon:
            if jid in demand[c["name"]]["jids"]:
                continue
            for a in c["aliases"]:
                if hits(a, jd_norm):
                    d = demand[c["name"]]
                    d["jids"].add(jid); d["titles"].append(title); d["jd_ids"].append(jid)
                    d["required"].add(jid)  # explicitly named in JD = required demand
                    d["phrases"].add(a)
                    break

    # ---- 3. Resume evidence per canonical skill (four levels) ----------------
    def resume_has_any(c):
        if has_resume and (any(hits(a, res) for a in c["aliases"]) or hits(c["name"], res)):
            return "direct"
        if has_resume and any(hits(a, res) for a in c.get("adjacent", [])):
            return "adjacent"
        return None

    # Historical proof: a skill the pasted analyses repeatedly said you "have"
    # is real resume evidence even when the live mirror is a stub.
    hist_have = {}  # name -> set(jids where you had it)
    for row in rows:
        try:
            rep = (json.loads(row["report"]) if isinstance(row["report"], str)
                   else row["report"]) or {}
        except Exception:
            rep = {}
        for h in (rep.get("have") or []):
            k = (h.get("key") or "").strip()
            if k in demand:
                hist_have.setdefault(k, set()).add(row["id"])

    # ---- 4. Assemble the unified skill list --------------------------------
    skills_out = []
    for c in canon:
        name = c["name"]
        d = demand[name]
        n_jd = len(d["jids"])
        n_req = len(d["required"])
        n_pref = len(d["preferred"])
        pct = round(100 * n_jd / jd_count) if jd_count else 0

        # Recruiter demand band from literal demand frequency.
        if n_jd == 0:
            demand_band = "—"
        elif pct >= 60:
            demand_band = "Very High"
        elif pct >= 40:
            demand_band = "High"
        elif pct >= 20:
            demand_band = "Medium"
        else:
            demand_band = "Low"

        # Resume evidence level (four levels, spec #8/#10).
        rh = resume_has_any(c)
        hist = hist_have.get(name, set())
        if rh == "direct" or (hist and len(hist) >= 2):
            # direct resume wording OR 2+ pasted analyses flagged it as
            # something you have -> demonstrated.
            level = "demonstrated"
        elif rh == "adjacent" or (hist and len(hist) == 1):
            # resume shows overlapping experience, or only weak historical
            # signal -> transferable/adjacent (never call it a gap just because
            # the exact keyword is absent, RULE 3).
            level = "transferable"
        elif n_jd > 0:
            # partial vs genuine gap: genuine gap only when demand is high AND
            # there is truly no resume or adjacent signal at all.
            level = "gap" if pct >= 40 else "partial"
        else:
            level = "na"

        # Market-vs-Learn action (spec #14).
        if level == "demonstrated":
            action = "MARKET MORE STRONGLY"
        elif level == "transferable":
            action = "STRENGTHEN / FORMALIZE"
        elif level == "partial":
            action = "STRENGTHEN / LEARN"
        elif level == "gap":
            action = "LEARN"
        else:
            action = "NOT PRIORITY"

        # Tier (spec #11, evidence-based; never Tier4 just for "not in resume").
        if pct >= 50 and (level == "demonstrated" or level == "transferable"):
            tier = 1
        elif pct >= 80 and level in ("gap", "partial"):
            tier = 1
        elif pct >= 40:
            tier = 2
        elif pct >= 20 or (n_jd > 0 and level in ("gap", "partial")):
            tier = 3
        else:
            tier = 4

        skills_out.append(dict(
            name=name, n_jd=n_jd, n_req=n_req, n_pref=n_pref, pct=pct,
            demand=demand_band, level=level, action=action, tier=tier,
            kind=c["kind"], category=c["category"], why=c["why"],
            titles=d["titles"], jd_ids=sorted(d["jd_ids"]),
            phrases=sorted(d["phrases"])[:6],
            hist_jds=sorted(hist)))

    # Sort: most-demanded first, then by name.
    skills_out.sort(key=lambda x: (-x["n_jd"], -x["pct"], x["name"]))

    # ---- 5. Tiers, Top 15, Core stack, Roadmap, etc. -----------------------
    tiers = {1: [], 2: [], 3: [], 4: []}
    for sk in skills_out:
        tiers[sk["tier"]].append(sk)

    top15 = [sk for sk in skills_out if sk["n_jd"] > 0][:15]

    # Core Career Skill Stack (spec #12) — populate only categories with demand.
    core_stack = {}
    for cat, member_names in CATEGORY_MEMBERS.items():
        demonstrated, transferable, needs = [], [], []
        members = [sk for sk in skills_out if sk["name"] in set(member_names)]
        if not members:
            continue
        for sk in members:
            if sk["level"] == "demonstrated":
                demonstrated.append(sk["name"])
            elif sk["level"] == "transferable":
                transferable.append(sk["name"])
            elif sk["level"] in ("partial", "gap") and sk["n_jd"] > 0:
                needs.append(sk["name"])
        if demonstrated or transferable or needs:
            core_stack[cat] = dict(demonstrated=demonstrated,
                                   transferable=transferable, needs=needs)

    # Learning roadmap (spec #13): demand x evidence x career relevance x gap.
    roadmap = dict(immediately=[], next=[], later=[], nopriority=[])
    for sk in skills_out:
        if sk["n_jd"] == 0:
            continue
        if sk["action"] in ("LEARN",) and sk["tier"] in (1, 2):
            roadmap["immediately"].append(sk)
        elif sk["action"] in ("LEARN", "STRENGTHEN / LEARN") and sk["tier"] == 3:
            roadmap["next"].append(sk)
        elif sk["action"] == "STRENGTHEN / FORMALIZE":
            roadmap["later"].append(sk)
        else:
            roadmap["nopriority"].append(sk)
    for k in roadmap:
        roadmap[k].sort(key=lambda x: -x["pct"])

    # Resume positioning (spec #16).
    market_more = [sk for sk in skills_out
                   if sk["level"] == "demonstrated" and sk["pct"] >= 40]
    genuine_gaps = [sk for sk in skills_out
                    if sk["level"] == "gap" and sk["n_jd"] > 0]
    tech_ai_gaps = [sk for sk in genuine_gaps
                    if any(w in sk["name"] for w in ("AI", "LMS", "Power BI",
                                                     "Excel", "CRM", "e-learning",
                                                     "analytics", "data"))]
    not_add = [sk for sk in skills_out
               if sk["level"] == "na" and sk["pct"] < 20]
    keywords = sorted(
        {(p, len([1 for row in rows if hits(p, norm(row["jd_text"] or ""))]))
         for sk in skills_out for p in sk["phrases"]},
        key=lambda x: -x[1])[:15]

    # Career direction (spec #15) — confidence scales with sample size.
    roles = {}
    for row in rows:
        role = row["role"] or "training manager"
        roles[role] = roles.get(role, 0) + 1
    dom_role = max(roles, key=roles.get) if roles else "training manager"
    confidence = ("PRELIMINARY — sample size too small for reliable conclusions"
                  if jd_count < 5 else
                  "EARLY TREND" if jd_count < 10 else
                  "EMERGING MARKET SIGNAL" if jd_count < 20 else
                  "MEANINGFUL CUMULATIVE SIGNAL")
    direction_spec = [
        ("Training Team Lead", "training manager"),
        ("Assistant Manager – Training", "training manager"),
        ("Assistant Manager – L&D", "l&d"),
        ("Training Manager", "training manager"),
        ("L&D Specialist", "l&d"),
        ("Sales Training Manager", "rsm"),
        ("Sales Enablement", "rsm"),
        ("Customer Experience / Customer Service Manager", "customer"),
    ]
    career_direction = []
    for role_name, fam in direction_spec:
        # Fit is driven by how many analyzed JDs are in that role family
        # (cumulative evidence) plus your demonstrated resume strength.
        fam_jds = sum(1 for row in rows if (row["role"] or "") == fam)
        if fam_jds >= 3:
            fit = "Strong match"
        elif fam_jds >= 1 and dom_role in (fam, "training manager", "l&d"):
            fit = "Possible — supported by cumulative JDs"
        else:
            fit = "Possible — limited JD evidence"
        career_direction.append(dict(role=role_name, fit=fit, fam_jds=fam_jds))

    # Final priority (spec #20).
    focus = [sk for sk in skills_out
             if sk["tier"] in (1, 2) and sk["n_jd"] > 0
             and sk["action"] in ("LEARN", "STRENGTHEN / LEARN",
                                  "STRENGTHEN / FORMALIZE")][:5]
    market_more_top = market_more[0]["name"] if market_more else None
    top_gap = (genuine_gaps[0]["name"] if genuine_gaps
               else next((sk["name"] for sk in skills_out
                          if sk["level"] in ("partial", "gap") and sk["n_jd"] > 0), None))
    ai_skill = (tech_ai_gaps[0]["name"] if tech_ai_gaps
                else next((sk["name"] for sk in skills_out
                           if "AI" in sk["name"] and sk["level"] in ("gap", "partial")), None))
    best_dir = next((d["role"] for d in career_direction if "Strong" in d["fit"]),
                    career_direction[0]["role"] if career_direction else "")

    if jd_count == 0:
        trend = "No job descriptions analyzed yet — paste a ChatGPT reply to begin."
    else:
        top_names = [sk["name"] for sk in top15[:3]]
        trend = (f"Across {jd_count} analyzed JDs, recruiters most repeatedly ask for: "
                 + ", ".join(top_names) + f". {confidence}.")

    return dict(
        jd_count=jd_count, small=small, roles=roles, dom_role=dom_role,
        skills=skills_out, tiers=tiers, tier_label=TIER_LABEL,
        level_label=LEVEL_LABEL, action_label=ACTION_LABEL,
        core_stack=core_stack, top15=top15, roadmap=roadmap,
        market_more=[sk["name"] for sk in market_more],
        genuine_gaps=[sk["name"] for sk in genuine_gaps],
        tech_ai_gaps=[sk["name"] for sk in tech_ai_gaps],
        not_add=not_add, keywords=keywords,
        career_direction=career_direction, confidence=confidence,
        final=dict(focus=[sk["name"] for sk in focus],
                   market_more=market_more_top, top_gap=top_gap,
                   ai_skill=ai_skill, best_dir=best_dir, trend=trend),
        generated=dt.datetime.now().strftime("%d %b %Y, %I:%M %p"))


# --------------------------------------------------------------- routes
@app.route("/cumulative")
def cumulative_page():
    """Read-only Career-level view: 'Cumulative Recruiter-Demand Analysis'.
    Computed from EVERY JD the user has analyzed, so it is real data and
    auto-updates whenever a new JD is added."""
    r = get_resume()
    if not r:
        flash("Save your resume first (Step 1 on the home page), then come back.", "bad")
        return redirect(url_for("home"))
    data = cumulative_analysis()
    return render_template("cumulative.html", c=data)


@app.route("/")
def home():
    import pasteback
    r = get_resume()
    prompt = ""
    if r:
        # Build the copy-ready prompt: resume already inside + the method.
        # JD is left as a placeholder the user fills in ChatGPT.
        prompt = pasteback.build_prompt(
            r["text"], "[PASTE THE JOB DESCRIPTION TEXT HERE — replace this whole line with the real job description before sending to ChatGPT]")
    with db() as c:
        jds = c.execute("SELECT id,title,role,created,report FROM jd ORDER BY id DESC LIMIT 15").fetchall()
    rows = []
    for j in jds:
        try:
            rr = json.loads(j["report"])
        except Exception:
            continue
        rows.append(dict(id=j["id"], title=j["title"], role=rr.get("role_label", ""),
                         created=j["created"], pct=rr.get("match_pct", 0), gaps=len(rr.get("gaps", []))))
    return render_template("home.html", resume=r, prompt=prompt, jds=rows,
                           gh_backup=reports_store.enabled(),
                           cum=cumulative_analysis())


@app.route("/resume", methods=["POST"])
def upload_resume():
    f = request.files.get("file")
    if not f or not f.filename:
        flash("Pick your resume file first.", "bad")
        return redirect(url_for("home"))
    name = (request.form.get("name") or "My resume").strip()[:40] or "My resume"
    text, note = extract_text(f)
    if len(text.strip()) < 50:
        flash(f"Could not read that resume ({note}). Try a PDF or DOCX.", "bad")
        return redirect(url_for("home"))
    path = os.path.join(UPLOADS, "resume_" + re.sub(r"[^A-Za-z0-9._-]", "_", f.filename))
    f.seek(0)
    with open(path, "wb") as out:
        out.write(f.read())
    set_resume(name, f.filename, text, dt.datetime.now().strftime("%d %b %Y, %I:%M %p"))
    flash(f"Resume '{name}' saved ({note}, {len(text.split())} words). It stays saved - no need to upload again.", "good")
    return redirect(url_for("home"))


@app.route("/delete-resume/<int:rid>", methods=["POST"])
def remove_resume(rid):
    delete_resume(rid)
    flash("Resume cleared.", "good")
    return redirect(url_for("home"))


@app.route("/delete-prompt/<int:pid>", methods=["POST"])
def remove_prompt(pid):
    delete_prompt(pid)
    flash("Saved prompt cleared.", "good")
    return redirect(url_for("home"))

@app.route("/report/<int:jd_id>")
def report(jd_id):
    with db() as c:
        row = c.execute("SELECT * FROM jd WHERE id=?", (jd_id,)).fetchone()
    if not row:
        flash("That report is gone.", "bad")
        return redirect(url_for("home"))
    # The Cumulative Recruiter-Demand Analysis is appended after every JD's
    # skill-gap report. It is computed server-side from ALL saved JDs so the
    # cross-JD frequencies are real (ChatGPT only sees the current JD).
    # We also compute the "before this JD" picture to report rank movements
    # (spec #10: tell the user when a new JD changed a skill's importance).
    cum = cumulative_analysis()
    rank_shift = []
    before = cumulative_analysis(exclude_id=jd_id)
    if before["jd_count"] >= 1:
        prev_rank = {sk["name"]: i + 1 for i, sk in enumerate(cum["top15"])}
        before_rank = {sk["name"]: i + 1 for i, sk in enumerate(before["top15"])}
        for sk in cum["top15"]:
            k = sk["name"]
            old = before_rank.get(k)
            if old and old != prev_rank[k]:
                rank_shift.append(dict(
                    key=k, old=old, new=prev_rank[k],
                    direction="up" if prev_rank[k] < old else "down",
                    pct=sk["pct"]))
    cum["rank_shift"] = rank_shift
    return render_template("report.html", jd=row, r=json.loads(row["report"]),
                           ROLE_EXPERIENCE=ROLE_EXPERIENCE, c=cum)


@app.route("/paste")
def paste_page():
    """Step 3: paste ChatGPT's reply back. The prompt lives on the home page
    (built with the resume already inside)."""
    r = get_resume()
    if not r:
        flash("Save your resume first (Step 1 on the home page), then come back here.", "bad")
        return redirect(url_for("home"))
    jd_prefill = ""
    jd_id = request.args.get("jd", type=int)
    if jd_id:
        with db() as c:
            jrow = c.execute("SELECT jd_text FROM jd WHERE id=?", (jd_id,)).fetchone()
        if jrow:
            jd_prefill = jrow["jd_text"] or ""
    return render_template("paste.html", jd_prefill=jd_prefill)


@app.route("/paste-back", methods=["POST"])
def paste_back():
    """Zero-API AI path, step 2: glue the AI's reply back into a full report.

    The JD is NOT sent to the app -- the user pastes it into ChatGPT alongside
    the app-generated prompt (resume + method). So here we only take the
    ChatGPT reply and let pasteback parse it into the structured report.
    Works for ANY role (the resume is already inside the prompt the user sent)."""
    import pasteback
    rid = request.form.get("resume_id", type=int)
    r = get_resume(rid)
    if not r:
        flash("Save your resume first (Step 1 on the home page), then come back.", "bad")
        return redirect(url_for("home"))
    reply = (request.form.get("reply") or "").strip()
    jd_text = (request.form.get("jd_text") or "").strip()  # optional, kept only for record
    title = (request.form.get("title") or "").strip() or "Pasted ChatGPT analysis"
    if len(reply) < 40:
        flash("Paste ChatGPT's full reply in the box before saving.", "bad")
        return redirect(url_for("paste_page"))
    rep, err = pasteback.from_paste(r["text"], jd_text, reply)
    if err:
        flash(err, "bad")
        return redirect(url_for("paste_page"))
    # STABLE ID: allocate a number that's > both the local DB max and the
    # GitHub max, so cold-boot re-seeds reuse the same ids and GitHub filenames
    # stay in sync with the DB. (If GitHub is off, this just falls back to the
    # local max + 1, identical to the old AUTOINCREMENT behaviour.)
    with db() as c:
        local_max = c.execute("SELECT COALESCE(MAX(id),0) FROM jd").fetchone()[0]
    new_id = max(local_max, reports_store.max_report_id()) + 1
    with db() as c:
        c.execute("INSERT INTO jd(id,title,role,source,jd_text,report,created) VALUES(?,?,?,?,?,?,?)",
                  (new_id, title, rep["role"], "paste-back (no API key)", jd_text, json.dumps(rep),
                   dt.datetime.now().strftime("%d %b %Y, %I:%M %p")))
    # Mirror to GitHub so this report survives Render's disk wipe on restart.
    if reports_store.enabled():
        reports_store.save_report(new_id, title, rep.get("role", ""),
                                 rep, jd_text,
                                 dt.datetime.now().strftime("%d %b %Y, %I:%M %p"))
    return redirect(url_for("report", jd_id=new_id))


@app.route("/delete/<int:jd_id>", methods=["POST"])
def delete(jd_id):
    with db() as c:
        c.execute("DELETE FROM jd WHERE id=?", (jd_id,))
    # Best-effort remove the durable copy too.
    if reports_store.enabled():
        reports_store.delete_report(jd_id)
    return redirect(url_for("home"))


def resume_draft(rep, jd_text, resume_text):
    """Build ready-to-paste resume wording tuned to this JD.

    Honesty rule: skills you already have -> reworded in JD language now.
    Skills you don't -> parked in a 'add this line the day you finish it' list,
    never silently claimed.
    """
    jd = norm(jd_text)
    role_label = rep["role_label"]

    # Naukri headline: role title + your real anchors + top JD keywords you own
    owned = [h["key"] for h in rep["have"]]
    headline = (f"{role_label} | 12-member trainer team | NHT & Technical Training owner | "
                + ", ".join(k.split(" (")[0] for k in owned[:3]))

    # Key Skills box on Naukri = exact JD keywords. Recruiter search runs on this.
    kw = []
    JUNK = {"team of", "coordinate a team", "od", "gt", "mt", "id", "pip", "sop", "ai"}
    for s in SKILLS:
        for a in s["aliases"]:
            if hits(a, jd) and len(a) > 3 and a.lower() not in JUNK:
                kw.append(a.title() if a.islower() else a)
    seen, keywords = set(), []
    for k in kw:
        if k.lower() not in seen:
            seen.add(k.lower())
            keywords.append(k)

    summary = (f"Training operations professional targeting {role_label} roles. "
               f"Own a 12-day New Hire Training curriculum and a 3-day technical training track "
               f"for a 12-member trainer team, with daily activity governance, MIS reporting and "
               f"audit-clean training records. Strengths this role asks for: "
               + "; ".join(k.split(" (")[0] for k in owned[:5]) + ".")

    # Rewrite existing strengths in the JD's own language
    now = [dict(skill=h["key"], bullet=RESUME_LINES.get(h["key"], h["proof"])) for h in rep["have"]]
    later = [dict(skill=g["key"], bullet=RESUME_LINES.get(g["key"], g["proof"]))
             for g in rep["gaps"] + rep["implied"]]

    return dict(headline=headline[:250], summary=summary,
                keywords=keywords[:25], now=now, later=later)


# Resume-ready phrasing per skill: what the bullet should literally say.
RESUME_LINES = {
    "ADDIE instructional design":
        "Designed and rebuilt the 12-day NHT curriculum using the ADDIE model - needs analysis, "
        "module design, content development, rollout and post-batch evaluation.",
    "Kirkpatrick training evaluation (L1-L4)":
        "Measured training effectiveness on Kirkpatrick L1-L4: feedback scores, assessment pass %, "
        "on-floor behaviour audits and business-metric impact reported to leadership.",
    "TNI / TNA (Training Needs Identification)":
        "Ran Training Needs Identification across a 12-trainer team using performance data, "
        "manager inputs and skill surveys; converted findings into a quarterly training calendar.",
    "LMS administration":
        "Administered the LMS end to end - course setup, batch enrolment, completion tracking "
        "and compliance/certification reporting.",
    "e-learning authoring (Articulate / Rise / Captivate)":
        "Built digital learning modules (Articulate Rise 360 / Storyline) converting classroom "
        "NHT content into self-paced e-learning with assessments.",
    "Train-the-Trainer / facilitation certification":
        "Certified Train-the-Trainer; facilitate classroom and virtual sessions and certify new "
        "trainers through structured demo evaluation.",
    "Coaching & feedback models (GROW, SBI)":
        "Coach trainers one-on-one using the GROW model and SBI feedback; improved "
        "underperformer output through structured monthly capability reviews.",
    "Content / SOP documentation":
        "Authored one-page SOPs and job aids for recurring processes (MIS Update, Attendance "
        "Regularization, Visit Barge / New Process), reducing repeat queries.",
    "Excel advanced (pivots, lookups, dashboards)":
        "Advanced Excel - pivot tables, XLOOKUP/VLOOKUP and slicer-driven dashboards used for "
        "daily trainer activity MIS across a 12-member team.",
    "Power BI / data visualisation":
        "Publish weekly Power BI dashboards on training throughput, attendance and certification "
        "status for management review.",
    "Learning analytics / MIS reporting":
        "Own training MIS and learning analytics - batch throughput, certification pass rate, "
        "TAT and attrition linked to training - reported in weekly and monthly reviews.",
    "Sales target ownership & pipeline management":
        "Carried and delivered a quarterly target of <fill in number>; managed pipeline coverage "
        "and conversion against plan.",
    "Channel / distributor management":
        "Managed channel/distributor partners across the territory - coverage planning, beat "
        "plans and partner ROI reviews.",
    "CRM tools (Salesforce / Zoho / LeadSquared)":
        "Hands-on with CRM (Salesforce/Zoho) for lead stages, activity logging and pipeline "
        "reporting; Salesforce Trailhead certified badges.",
    "Team leadership & performance management":
        "Coordinate a 12-member trainer team - rostering, daily activity governance, capability "
        "reviews and performance conversations.",
    "Stakeholder / business partnering":
        "Partner with business stakeholders and senior management; own the weekly/monthly training "
        "review deck and governance cadence.",
    "Onboarding / new hire training programme design":
        "Own the 12-day New Hire Training programme and 3-day Technical Training track - "
        "objectives, assessments, certification gates and ramp-up tracking.",
    "Compliance / audit readiness of training records":
        "Maintain audit-ready training records and attendance regularization; closed audit cycles "
        "with zero documentation observations.",
    "Communication & business English / presentation":
        "Strong classroom facilitation, business communication and presentation skills; deliver "
        "leadership-facing review presentations.",
    "AI tools for L&D (ChatGPT/Gemini for content)":
        "Use AI tools (ChatGPT/Gemini/Copilot) to draft assessments, role-play scripts and "
        "training content, cutting content development time significantly.",
    "Vendor & training budget management":
        "Manage training budget and external vendor partners; track cost per trainee against plan.",
}


@app.route("/resume-draft/<int:jd_id>")
def resume_draft_page(jd_id):
    with db() as c:
        row = c.execute("SELECT * FROM jd WHERE id=?", (jd_id,)).fetchone()
    r = get_resume()
    if not row or not r:
        flash("Need a saved resume and a job first.", "bad")
        return redirect(url_for("home"))
    rep = json.loads(row["report"])
    return render_template("draft.html", jd=row, r=rep,
                           d=resume_draft(rep, row["jd_text"], r["text"]))


@app.route("/icon-192.png")
def icon192():
    from flask import send_from_directory
    return send_from_directory(os.path.join(BASE, "static"), "icon-192.png")


@app.route("/icon-512.png")
def icon512():
    from flask import send_from_directory
    return send_from_directory(os.path.join(BASE, "static"), "icon-512.png")


def lan_ip():
    """Best-guess LAN IP so the phone can reach this laptop."""
    import socket as sk
    try:
        s = sk.socket(sk.AF_INET, sk.SOCK_DGRAM)
        s.connect(("8.8.8.8", 80))
        ip = s.getsockname()[0]
        s.close()
        return ip
    except Exception:
        return "127.0.0.1"


@app.route("/qr.png")
def qr_png():
    """QR code of the phone URL so you can just scan it."""
    import qrcode
    from flask import Response
    img = qrcode.make(f"http://{lan_ip()}:5055/")
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return Response(buf.getvalue(), mimetype="image/png")




@app.route("/backup")
def backup():
    """Download everything (all saved resumes + all reports) as one JSON file.

    Needed because Render's free plan erases files when the app redeploys or
    restarts. Keep the downloaded file; restore it below in one click.
    """
    from flask import Response
    with db() as c:
        rrows = c.execute("SELECT name,filename,text,uploaded FROM resumes").fetchall()
        jds = c.execute("SELECT title,role,source,jd_text,report,created FROM jd").fetchall()
    blob = dict(resumes=[dict(x) for x in rrows], jobs=[dict(x) for x in jds],
                saved=dt.datetime.now().isoformat())
    return Response(json.dumps(blob, indent=1),
                    mimetype="application/json",
                    headers={"Content-Disposition":
                             "attachment; filename=career-coach-backup.json"})


@app.route("/restore", methods=["POST"])
def restore():
    f = request.files.get("bak")
    if not f or not f.filename:
        flash("Pick your backup file first.", "bad")
        return redirect(url_for("home"))
    try:
        blob = json.loads(f.read().decode("utf-8", "ignore"))
    except Exception:
        flash("That file isn't a valid backup.", "bad")
        return redirect(url_for("home"))
    resumes = list(blob.get("resumes") or [])
    if blob.get("resume"):  # legacy single-resume backup shape
        resumes.append(blob["resume"])
    jobs = blob.get("jobs", [])
    with db() as c:
        for b in resumes:
            c.execute("INSERT INTO resumes(name,filename,text,uploaded) VALUES(?,?,?,?)",
                      (b.get("name", "My resume"), b.get("filename"), b.get("text"),
                       b.get("uploaded")))
        # Mirror restored resume to GitHub too (so it survives future restarts).
        if resumes and reports_store.enabled():
            b = resumes[-1]
            reports_store.save_resume(b.get("name", "My resume"), b.get("filename"),
                                      b.get("text"), b.get("uploaded"))
        # Allocate stable ids aligned with any GitHub copies already present.
        base = max(reports_store.max_report_id(),
                   c.execute("SELECT COALESCE(MAX(id),0) FROM jd").fetchone()[0])
        for n, j in enumerate(jobs, start=1):
            jid = base + n
            rep = None
            try:
                rep = json.loads(j.get("report")) if isinstance(j.get("report"), str) else j.get("report")
            except Exception:
                rep = j.get("report")
            c.execute("INSERT INTO jd(id,title,role,source,jd_text,report,created) VALUES(?,?,?,?,?,?,?)",
                      (jid, j.get("title"), j.get("role"), j.get("source"), j.get("jd_text"),
                       json.dumps(rep) if rep is not None else "{}", j.get("created")))
            if reports_store.enabled() and rep is not None:
                reports_store.save_report(jid, j.get("title"), j.get("role"), rep,
                                          j.get("jd_text", "") or "", j.get("created", ""))
    flash(f"Restored {len(resumes)} resume(s) and {len(jobs)} job(s).", "good")
    return redirect(url_for("home"))




if __name__ == "__main__":
    init_db()
    port = int(os.environ.get("PORT", 5055))
    print(f"\n  Career Coach running -> http://127.0.0.1:{port}\n")
    app.run(host="0.0.0.0", port=port, debug=False)
